"""PII redaction — deterministic detection + format-preserving rewriting.

Covers the detector (SSN/ITIN/EIN, email, phone, Luhn-gated cards, labeled
accounts, DOB, addresses, extra_terms, overlap resolution), the three styles,
per-format redactors (docx styling survives; xlsx formulas untouched; PDF
rebuilt with content truly gone), and the ``redact_pii`` tool contract
(original never modified, workspace-confined output, honest counts).
"""

from __future__ import annotations

from pathlib import Path

from iron_jarvis.documents import extract_text, write_document
from iron_jarvis.documents.redact import (
    find_pii_spans,
    mask_text,
    redact_file,
)
from iron_jarvis.documents.tools import document_tools
from iron_jarvis.tools.base import ToolContext

SSN_TEXT = "Taxpayer SSN 123-45-6789 filed with spouse (ssn: 987654321)."


def _cats(text, **kw):
    return [c for _s, _e, c in find_pii_spans(text, **kw)]


# --- detection ---------------------------------------------------------------


def test_detects_hyphenated_and_labeled_ssn():
    cats = _cats(SSN_TEXT)
    assert "ssn" in cats and "ssn_labeled" in cats


def test_detects_itin_ein_email_phone():
    text = (
        "ITIN 912-93-1234, EIN 12-3456789, mail bruce@wayne.example.com, "
        "call (212) 555-0187."
    )
    cats = _cats(text)
    assert {"itin", "ein", "email", "phone"} <= set(cats)


def test_credit_card_requires_luhn():
    assert "credit_card" in _cats("card 4539 1488 0343 6467")  # Luhn-valid
    assert "credit_card" not in _cats("ref 4539 1488 0343 6468")  # invalid check digit


def test_context_gated_account_and_dob():
    cats = _cats("Account #: 001234567 · DOB: 04/15/1980")
    assert "bank_account" in cats and "dob" in cats
    # Bare digits / bare dates never match without their context words.
    assert _cats("invoice 001234567 dated 04/15/1980") == []


def test_street_address_detected():
    assert "address" in _cats("Mail to 1234 Evergreen Terrace, Springfield")


def test_extra_terms_case_insensitive_and_short_terms_ignored():
    cats = _cats("Prepared for John Doe (JOHN DOE).", extra_terms=["john doe", "J"])
    assert cats == ["custom", "custom"]


def test_categories_filter_limits_detection():
    text = "SSN 123-45-6789, email a@b.example.com"
    cats = _cats(text, categories={"email"})
    assert cats == ["email"]


# --- masking styles ----------------------------------------------------------


def test_black_style_preserves_length():
    masked, counts = mask_text("SSN 123-45-6789 end", style="black")
    assert "123-45-6789" not in masked
    assert "█" * 11 in masked
    assert len(masked) == len("SSN 123-45-6789 end")
    assert counts == {"ssn": 1}


def test_label_style_tags_category():
    masked, _ = mask_text("reach me at clark@daily.example.com", style="label")
    assert "[EMAIL]" in masked and "@" not in masked


def test_remove_style_deletes_value():
    masked, _ = mask_text("EIN 12-3456789.", style="remove")
    assert masked == "EIN ."


# --- format redactors --------------------------------------------------------


def test_docx_redaction_preserves_styling_and_original(tmp_path):
    import docx

    src, dst = tmp_path / "letter.docx", tmp_path / "letter.redacted.docx"
    d = docx.Document()
    par = d.add_paragraph()
    par.add_run("Client SSN ").bold = True
    par.add_run("123-45-").bold = True  # PII split across runs on purpose
    par.add_run("6789 attached.")
    d.add_table(rows=1, cols=1).cell(0, 0).text = "EIN 12-3456789"
    d.save(str(src))
    before = src.read_bytes()

    counts, _note = redact_file(src, dst, style="black")
    assert counts.get("ssn") == 1 and counts.get("ein") == 1
    assert src.read_bytes() == before  # original untouched

    out = docx.Document(str(dst))
    text = "\n".join(p.text for p in out.paragraphs)
    assert "123-45" not in text and "█" in out.paragraphs[0].text
    assert out.paragraphs[0].runs[0].bold is True  # styling survived
    assert "12-3456789" not in out.tables[0].cell(0, 0).text


def test_xlsx_redaction_keeps_formulas_and_numbers(tmp_path):
    from openpyxl import Workbook, load_workbook

    src, dst = tmp_path / "book.xlsx", tmp_path / "book.redacted.xlsx"
    wb = Workbook()
    ws = wb.active
    ws["A1"] = "SSN 123-45-6789"
    ws["A2"] = 123456789  # a NUMBER — must never be rewritten
    ws["A3"] = "=SUM(1,2)"
    wb.save(str(src))

    counts, note = redact_file(src, dst, style="label")
    assert counts == {"ssn": 1}
    out = load_workbook(str(dst))
    ws2 = out.active
    assert "[SSN]" in ws2["A1"].value
    assert ws2["A2"].value == 123456789
    assert ws2["A3"].value == "=SUM(1,2)"
    assert "formulas" in note


def test_pdf_rebuilt_with_content_truly_gone(tmp_path):
    src, dst = tmp_path / "return.pdf", tmp_path / "return.redacted.pdf"
    write_document(src, "Taxpayer 123-45-6789 owes nothing.")
    counts, note = redact_file(src, dst, style="black")
    assert counts.get("ssn") == 1
    assert "rebuilt" in note.lower()
    assert "123-45-6789" not in extract_text(dst)  # truly gone, not painted over


def test_csv_and_txt_redaction(tmp_path):
    src, dst = tmp_path / "list.csv", tmp_path / "list.redacted.csv"
    src.write_text("name,ssn\nJane,123-45-6789\n", encoding="utf-8")
    counts, _ = redact_file(src, dst, style="remove", extra_terms=["Jane"])
    assert counts == {"ssn": 1, "custom": 1}
    out = dst.read_text(encoding="utf-8")
    assert "123-45-6789" not in out and "Jane" not in out
    assert out.startswith("name,ssn")  # everything else intact


# --- the tool ----------------------------------------------------------------


def _ctx(workspace: Path) -> ToolContext:
    return ToolContext(
        workspace=workspace,
        session_id="t",
        agent_run_id="t",
        config=None,
        event_bus=None,
        engine=None,
    )


def _tool():
    return next(t for t in document_tools() if t.name == "redact_pii")


async def test_tool_default_output_and_untouched_source(tmp_path):
    ws = tmp_path / "ws"
    ws.mkdir()
    src = ws / "client.txt"
    src.write_text(SSN_TEXT, encoding="utf-8")
    before = src.read_bytes()

    res = await _tool().execute({"path": "client.txt"}, _ctx(ws))
    assert res.ok, res.error
    assert res.data["path"] == "client.redacted.txt"
    assert res.data["total"] == 2
    assert src.read_bytes() == before
    assert "123-45-6789" not in (ws / "client.redacted.txt").read_text(encoding="utf-8")
    # The output line reports counts, never values.
    assert "123-45-6789" not in res.output


async def test_tool_outside_source_lands_in_workspace(tmp_path):
    outside = tmp_path / "elsewhere" / "notes.txt"
    outside.parent.mkdir(parents=True)
    outside.write_text("call (212) 555-0187", encoding="utf-8")
    ws = tmp_path / "ws"
    ws.mkdir()

    res = await _tool().execute({"path": str(outside)}, _ctx(ws))
    assert res.ok, res.error
    assert res.data["path"] == "notes.redacted.txt"
    assert (ws / "notes.redacted.txt").is_file()


async def test_tool_rejects_bad_style_and_category(tmp_path):
    ws = tmp_path / "ws"
    ws.mkdir()
    (ws / "a.txt").write_text("x", encoding="utf-8")
    bad_style = await _tool().execute({"path": "a.txt", "style": "invisible"}, _ctx(ws))
    assert not bad_style.ok and "style" in (bad_style.error or "")
    bad_cat = await _tool().execute({"path": "a.txt", "categories": ["magic"]}, _ctx(ws))
    assert not bad_cat.ok and "categories" in (bad_cat.error or "")


async def test_tool_refuses_overwriting_the_source(tmp_path):
    ws = tmp_path / "ws"
    ws.mkdir()
    (ws / "a.txt").write_text("SSN 123-45-6789", encoding="utf-8")
    res = await _tool().execute({"path": "a.txt", "output_path": "a.txt"}, _ctx(ws))
    assert not res.ok and "never overwritten" in (res.error or "")


async def test_tool_no_pii_is_honest(tmp_path):
    ws = tmp_path / "ws"
    ws.mkdir()
    (ws / "clean.txt").write_text("nothing sensitive here", encoding="utf-8")
    res = await _tool().execute({"path": "clean.txt"}, _ctx(ws))
    assert res.ok
    assert res.data["total"] == 0
    assert "no PII found" in res.output
