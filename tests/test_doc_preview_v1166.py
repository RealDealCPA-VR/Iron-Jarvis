"""v1.166.0 P1: docs-preview backend.

* /documents/preview: images return a POINTER (`kind:"image"`, no bytes in
  JSON); .csv/.tsv render as `kind:"sheet"` via stdlib csv with the same caps
  as xlsx (80 rows, 30 cols, 80-char cells) and an honest real extent.
* /documents/file: `?download=1` forces attachment; otherwise PDFs and images
  serve INLINE with a real content-type, everything else stays attachment.
* Truncation honesty: sheet payloads carry `total_rows`/`total_cols`;
  text/markdown/html payloads carry `total_chars` (length BEFORE the clip).

Every existing response field is preserved — these tests assert the old
fields alongside the new ones. Offline throughout.
"""

from __future__ import annotations

from fastapi.testclient import TestClient
from openpyxl import Workbook

from iron_jarvis.daemon.app import create_app


def _ledger(path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Ledger"
    ws.append(["Client", "Amount", "Month"])
    ws.append(["Acme", 1200, "Jan"])
    ws.append(["Birch", 800, "Jan"])
    ws.append(["Acme", 300, "Feb"])
    ws.append(["Cedar", "n/a", "Feb"])
    wb.save(str(path))


# --- image previews -----------------------------------------------------------


def test_preview_image_is_a_pointer_not_bytes(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    img = tmp_path / "photo.PNG"  # uppercase suffix must match too
    img.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 32)
    out = client.get("/documents/preview", params={"path": str(img)}).json()
    assert out["kind"] == "image"
    assert out["path"] == str(img)
    assert out["name"] == "photo.PNG" and out["suffix"] == ".png"
    # A pointer only — no payload fields carrying bytes or text.
    assert "content" not in out and "html" not in out and "rows" not in out


def test_preview_image_keeps_the_path_gate(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    assert client.get(
        "/documents/preview", params={"path": "relative.png"}
    ).status_code == 400
    assert client.get(
        "/documents/preview", params={"path": str(tmp_path / "ghost.png")}
    ).status_code == 404


# --- csv / tsv as sheets ------------------------------------------------------


def test_preview_csv_is_a_sheet_with_real_extent(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    f = tmp_path / "clients.csv"
    f.write_text("Client,Amount\nAcme,1200\nBirch,800\n", encoding="utf-8")
    out = client.get("/documents/preview", params={"path": str(f)}).json()
    assert out["kind"] == "sheet"
    assert out["sheets"] == ["CSV"] and out["sheet"] == "CSV"
    assert out["rows"] == [["Client", "Amount"], ["Acme", "1200"], ["Birch", "800"]]
    assert out["truncated"] is False
    assert out["total_rows"] == 3 and out["total_cols"] == 2


def test_preview_tsv_splits_on_tabs_not_commas(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    f = tmp_path / "data.tsv"
    f.write_text("a,x\tb\nc\td,y\n", encoding="utf-8")
    out = client.get("/documents/preview", params={"path": str(f)}).json()
    assert out["kind"] == "sheet"
    # Commas stay INSIDE cells — only tabs delimit.
    assert out["rows"] == [["a,x", "b"], ["c", "d,y"]]
    assert out["total_rows"] == 2 and out["total_cols"] == 2


def test_preview_csv_caps_match_xlsx_and_report_honestly(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    f = tmp_path / "wide.csv"
    lines = [",".join(f"r{i}c{j}" for j in range(35)) for i in range(120)]
    f.write_text("\n".join(lines) + "\n", encoding="utf-8")
    out = client.get("/documents/preview", params={"path": str(f)}).json()
    assert out["kind"] == "sheet"
    assert len(out["rows"]) == 80  # row cap
    assert len(out["rows"][0]) == 30  # col cap
    assert out["rows"][79][0] == "r79c0"  # the first 80 rows, in order
    assert out["truncated"] is True
    # The REAL extent, not the shown extent.
    assert out["total_rows"] == 120 and out["total_cols"] == 35


def test_preview_csv_clips_cells_at_80_chars(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    f = tmp_path / "long.csv"
    f.write_text("short," + "x" * 200 + "\n", encoding="utf-8")
    out = client.get("/documents/preview", params={"path": str(f)}).json()
    assert out["rows"][0][0] == "short"
    assert out["rows"][0][1] == "x" * 80  # clipped, same cap as xlsx cells


def test_preview_csv_replaces_bad_utf8_instead_of_erroring(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    f = tmp_path / "latin.csv"
    f.write_bytes(b"a,\xffb\n")  # invalid utf-8 byte
    r = client.get("/documents/preview", params={"path": str(f)})
    assert r.status_code == 200
    assert r.json()["rows"] == [["a", "�b"]]


def test_preview_csv_oversized_field_is_422_not_500(tmp_path):
    # A field over csv.field_size_limit (~128KB) makes csv.reader raise
    # _csv.Error, not OSError — a realistic export with a blob cell. That
    # must surface as the same honest 422 every sibling branch returns,
    # never an anonymous 500. (The limit itself is process-global state and
    # must NOT be raised to accommodate the file.)
    client = TestClient(create_app(str(tmp_path)))
    f = tmp_path / "blob.csv"
    f.write_text("id,payload\n1," + "x" * 200_000 + "\n", encoding="utf-8")
    r = client.get("/documents/preview", params={"path": str(f)})
    assert r.status_code == 422
    detail = r.json()["detail"]
    assert detail.startswith("could not read:")
    assert "field larger" in detail  # the csv error text reaches the user


# --- truncation honesty on existing kinds -------------------------------------


def test_preview_xlsx_reports_total_rows_and_cols(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    book = tmp_path / "book.xlsx"
    _ledger(book)
    out = client.get("/documents/preview", params={"path": str(book)}).json()
    # Every pre-v1.166.0 field is still there…
    assert out["kind"] == "sheet" and out["sheets"] == ["Ledger"]
    assert out["rows"][0] == ["Client", "Amount", "Month"]
    assert out["truncated"] is False
    # …and the real extent was added.
    assert out["total_rows"] == 5 and out["total_cols"] == 3


def test_preview_text_reports_total_chars(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    big = tmp_path / "big.txt"
    big.write_text("x" * 25_000, encoding="utf-8")
    out = client.get("/documents/preview", params={"path": str(big)}).json()
    assert out["kind"] == "text"
    assert len(out["content"]) == 20_000 and out["truncated"] is True
    assert out["total_chars"] == 25_000  # the length BEFORE the clip
    small = tmp_path / "small.txt"
    small.write_text("only forty-one characters in this file :)", encoding="utf-8")
    out2 = client.get("/documents/preview", params={"path": str(small)}).json()
    assert out2["truncated"] is False
    assert out2["total_chars"] == 41 == len(out2["content"])


def test_preview_markdown_reports_total_chars(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    note = tmp_path / "note.md"
    note.write_text("# Hello\nbody", encoding="utf-8")
    out = client.get("/documents/preview", params={"path": str(note)}).json()
    assert out["kind"] == "markdown" and "# Hello" in out["content"]
    assert out["truncated"] is False
    # total_chars measures the EXTRACTED text (extract_text may normalize a
    # trailing newline) — unclipped, it must equal what was returned.
    assert out["total_chars"] == len(out["content"]) >= len("# Hello\nbody")


def test_preview_docx_html_reports_total_chars(tmp_path):
    from docx import Document

    doc = Document()
    doc.add_heading("Engagement Letter", level=1)
    doc.add_paragraph("Dear client, thank you.")
    target = tmp_path / "letter.docx"
    doc.save(str(target))
    client = TestClient(create_app(str(tmp_path)))
    out = client.get("/documents/preview", params={"path": str(target)}).json()
    assert out["kind"] == "html" and "Engagement Letter" in out["html"]
    assert out["truncated"] is False
    assert out["total_chars"] == len(out["html"])  # unclipped ⇒ exact match


# --- /documents/file disposition + content-type -------------------------------


def test_file_pdf_serves_inline_as_pdf(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    f = tmp_path / "doc.pdf"
    f.write_bytes(b"%PDF-1.4 fake body")
    r = client.get("/documents/file", params={"path": str(f)})
    assert r.status_code == 200
    assert r.headers["content-type"] == "application/pdf"
    assert r.headers["content-disposition"].startswith("inline")
    assert 'filename="doc.pdf"' in r.headers["content-disposition"]


def test_file_png_serves_inline_as_image(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    f = tmp_path / "chart.png"
    f.write_bytes(b"\x89PNG\r\n\x1a\n" + b"\x00" * 8)
    r = client.get("/documents/file", params={"path": str(f)})
    assert r.status_code == 200
    assert r.headers["content-type"] == "image/png"
    assert r.headers["content-disposition"].startswith("inline")
    assert b"\x89PNG" in r.content  # the actual bytes, not a JSON pointer


def test_file_download_param_forces_attachment(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    pdf = tmp_path / "doc.pdf"
    pdf.write_bytes(b"%PDF-1.4 fake body")
    r = client.get("/documents/file", params={"path": str(pdf), "download": 1})
    assert r.headers["content-disposition"].startswith("attachment")
    assert r.headers["content-type"] == "application/pdf"  # type stays honest
    png = tmp_path / "chart.png"
    png.write_bytes(b"\x89PNG fake")
    r2 = client.get("/documents/file", params={"path": str(png), "download": 1})
    assert r2.headers["content-disposition"].startswith("attachment")


def test_file_non_renderable_kinds_stay_attachment(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    book = tmp_path / "book.xlsx"
    _ledger(book)
    r = client.get("/documents/file", params={"path": str(book)})
    assert r.headers["content-type"] == "application/octet-stream"
    assert r.headers["content-disposition"].startswith("attachment")
    txt = tmp_path / "note.txt"
    txt.write_text("plain", encoding="utf-8")
    r2 = client.get("/documents/file", params={"path": str(txt)})
    assert r2.headers["content-disposition"].startswith("attachment")


def test_file_gate_unchanged_with_new_param(tmp_path):
    client = TestClient(create_app(str(tmp_path)))
    assert client.get(
        "/documents/file", params={"path": "rel.pdf", "download": 1}
    ).status_code == 400
    assert client.get(
        "/documents/file", params={"path": str(tmp_path / "ghost.pdf")}
    ).status_code == 404


def test_xlsx_wide_sheet_column_clip_sets_truncated(tmp_path):
    """v1.167.0: a 40-column ledger must not render as a complete-looking
    30-column table — the column drop is truncation and must say so."""
    from openpyxl import Workbook

    p = tmp_path / "wide.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.append([f"col{i}" for i in range(40)])  # 40 cols, 1 row — no row clip
    wb.save(p)
    client = TestClient(create_app(str(tmp_path)))
    r = client.get("/documents/preview", params={"path": str(p)})
    assert r.status_code == 200
    body = r.json()
    assert body["truncated"] is True, "column clip was silent"
    assert len(body["rows"][0]) == 30
    assert body.get("total_cols") == 40


def test_xlsx_long_cell_clip_sets_truncated(tmp_path):
    from openpyxl import Workbook

    p = tmp_path / "longcell.xlsx"
    wb = Workbook()
    wb.active.append(["x" * 200])  # one cell over the 80-char clip
    wb.save(p)
    client = TestClient(create_app(str(tmp_path)))
    body = client.get("/documents/preview", params={"path": str(p)}).json()
    assert body["truncated"] is True, "cell clip was silent"
    assert body["rows"][0][0] == "x" * 80


def test_xlsx_small_sheet_stays_untruncated(tmp_path):
    from openpyxl import Workbook

    p = tmp_path / "small.xlsx"
    wb = Workbook()
    wb.active.append(["a", "b"])
    wb.save(p)
    client = TestClient(create_app(str(tmp_path)))
    body = client.get("/documents/preview", params={"path": str(p)}).json()
    assert body["truncated"] is False  # honesty cuts both ways


def test_csv_long_cell_clip_sets_truncated(tmp_path):
    p = tmp_path / "longcell.csv"
    p.write_text("short," + "y" * 200 + "\n", encoding="utf-8")
    client = TestClient(create_app(str(tmp_path)))
    body = client.get("/documents/preview", params={"path": str(p)}).json()
    assert body["truncated"] is True, "csv cell clip was silent"
    assert body["rows"][0][1] == "y" * 80
