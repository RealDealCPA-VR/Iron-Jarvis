"""Self-review QA loop (v1.134.0): deterministic document lint + one bounded
refinement round.

The lint half is pinned against REAL defective files built with
python-docx/openpyxl directly — every finding code has a seeded trigger, and
clean documents must produce zero findings (a noisy linter would burn LLM
refinement rounds on healthy deliverables). The refinement half is pinned via
the scripted fake router: a corrected reply replaces the deliverable, garbage
or a provider failure keeps the original with the error RECORDED, and — the
keep-better rule — a refinement that lints the same or worse never lands.
The code-appended "not included" honesty sections must survive refinement.
"""

from __future__ import annotations

import json
from pathlib import Path

import docx as docx_mod
from openpyxl import Workbook, load_workbook

from iron_jarvis.documents import batch as batch_mod
from iron_jarvis.documents.batch import refine_deliverable, run_batch
from iron_jarvis.documents.lint import (
    TRUNCATION_MIN_DOC_CHARS,
    lint_docx,
    lint_document,
    lint_xlsx,
)
from iron_jarvis.documents.readers import extract_text
from iron_jarvis.documents.tools import document_tools
from iron_jarvis.providers.adapters.base import LLMResponse
from iron_jarvis.providers.router import RouteResult
from iron_jarvis.tools.base import ToolContext


class ScriptedRouter:
    """Replays scripted completion texts in order; records every prompt.

    An ``Exception`` item raises instead (a real-provider failure)."""

    def __init__(self, replies, provider="anthropic"):
        self.replies = list(replies)
        self.provider = provider
        self.calls: list[tuple[str, str]] = []  # (system, user content)

    async def complete(self, *, system, messages, tools, task_class=None, **kw):
        assert len(messages) == 1 and messages[0].role == "user"
        self.calls.append((system, messages[0].content))
        assert self.replies, "router called more times than the script allows"
        item = self.replies.pop(0)
        if isinstance(item, Exception):
            raise item
        return RouteResult(LLMResponse(text=item), self.provider, "test-model")


def _ext_reply(summary, facts=()):
    return json.dumps(
        {
            "summary": summary,
            "facts": list(facts),
            "entities": {"people": [], "orgs": [], "dates": [], "amounts": []},
            "figures": [],
        }
    )


def _docs(folder: Path, **bodies: str) -> None:
    folder.mkdir(parents=True, exist_ok=True)
    for name, body in bodies.items():
        (folder / name.replace("_", ".")).write_text(body, encoding="utf-8")


def _ctx(ws: Path) -> ToolContext:
    ws.mkdir(parents=True, exist_ok=True)
    return ToolContext(
        workspace=ws, session_id="t", agent_run_id="t",
        config=None, event_bus=None, engine=None,
    )


def _codes(res) -> list[str]:
    return [f["code"] for f in res["findings"]]


def _check_shape(res) -> None:
    """The finding contract every consumer (batch qa, tool output) relies on."""
    assert set(res) == {"findings", "ok"}
    for f in res["findings"]:
        assert set(f) == {"code", "where", "detail", "severity"}
        assert f["severity"] in ("warn", "error")
        assert f["code"] and f["where"] and f["detail"]
    assert res["ok"] == (not any(f["severity"] == "error" for f in res["findings"]))


# ------------------------------------------------------------- docx lint -----


def test_docx_empty_document(tmp_path):
    d = docx_mod.Document()
    d.add_paragraph("   ")  # whitespace-only is still empty
    p = tmp_path / "empty.docx"
    d.save(str(p))
    res = lint_docx(p)
    _check_shape(res)
    assert _codes(res) == ["docx-empty"] and not res["ok"]


def test_docx_empty_sections(tmp_path):
    d = docx_mod.Document()
    d.add_heading("Intro", level=1)      # empty: next same-level heading follows
    d.add_heading("Data", level=1)
    d.add_paragraph("real body content")
    d.add_heading("End", level=1)        # empty: trailing heading, nothing after
    p = tmp_path / "sections.docx"
    d.save(str(p))
    res = lint_docx(p)
    _check_shape(res)
    assert _codes(res) == ["docx-empty-section", "docx-empty-section"]
    assert not res["ok"]
    wheres = [f["where"] for f in res["findings"]]
    assert any("Intro" in w for w in wheres) and any("End" in w for w in wheres)
    assert not any("Data" in w for w in wheres)


def test_docx_subheading_with_body_saves_the_parent_section(tmp_path):
    # "# A" followed only by "## B" + text: the text IS body for both — deeper
    # subheads neither end the parent section nor count as its body themselves.
    d = docx_mod.Document()
    d.add_heading("A", level=1)
    d.add_heading("B", level=2)
    d.add_paragraph("text under B")
    p = tmp_path / "nested.docx"
    d.save(str(p))
    assert lint_docx(p)["findings"] == []


def test_docx_zero_data_row_table_and_table_counts_as_body(tmp_path):
    d = docx_mod.Document()
    d.add_heading("Tables", level=1)
    t = d.add_table(rows=1, cols=2)  # header row only — no data
    t.cell(0, 0).text = "Name"
    t.cell(0, 1).text = "Total"
    p = tmp_path / "table.docx"
    d.save(str(p))
    res = lint_docx(p)
    _check_shape(res)
    # The table is flagged, but it COUNTS as body — no empty-section finding.
    assert _codes(res) == ["docx-table-no-data"] and not res["ok"]
    assert "table 1" in res["findings"][0]["where"]

    d2 = docx_mod.Document()
    d2.add_heading("Tables", level=1)
    t2 = d2.add_table(rows=2, cols=2)
    t2.cell(0, 0).text = "Name"
    t2.cell(1, 0).text = "row"
    p2 = tmp_path / "table_ok.docx"
    d2.save(str(p2))
    assert lint_docx(p2)["findings"] == []


def test_docx_numbering_gap(tmp_path):
    d = docx_mod.Document()
    d.add_paragraph("Steps:")
    d.add_paragraph("1. alpha")
    d.add_paragraph("3. charlie")  # "2." is missing
    p = tmp_path / "gap.docx"
    d.save(str(p))
    res = lint_docx(p)
    _check_shape(res)
    assert _codes(res) == ["docx-numbering-gap"]
    assert res["ok"]  # warn severity — numbering gaps alone don't refine
    assert '"1."' in res["findings"][0]["detail"]
    assert '"3."' in res["findings"][0]["detail"]

    d2 = docx_mod.Document()
    for line in ("1. a", "2. b", "3. c", "prose break", "1. restart", "2. fine"):
        d2.add_paragraph(line)
    p2 = tmp_path / "nogap.docx"
    d2.save(str(p2))
    assert lint_docx(p2)["findings"] == []  # sequences and restarts never flag


def _sized_doc(tail: str, *, filler_paras: int) -> docx_mod.Document:
    d = docx_mod.Document()
    d.add_heading("Report", level=1)
    for _ in range(filler_paras):
        d.add_paragraph(("filler sentence with content. " * 25).strip())
    d.add_paragraph(tail)
    return d


_DANGLING = "The committee concluded that the revenue outlook remains contin"


def test_docx_truncated_tail_fires_only_past_the_size_boundary(tmp_path):
    # Large doc (>= TRUNCATION_MIN_DOC_CHARS) ending mid-sentence: flagged.
    p = tmp_path / "big.docx"
    _sized_doc(_DANGLING, filler_paras=10).save(str(p))
    res = lint_docx(p)
    _check_shape(res)
    assert _codes(res) == ["docx-truncated-tail"]
    assert res["ok"]  # conservative heuristic -> warn, never error

    # Same ending on a SMALL doc: a style choice, not a clipped reply.
    small = tmp_path / "small.docx"
    _sized_doc(_DANGLING, filler_paras=1).save(str(small))
    assert lint_docx(small)["findings"] == []

    # Large doc ending with terminal punctuation: clean.
    done = tmp_path / "done.docx"
    _sized_doc(_DANGLING + "gent.", filler_paras=10).save(str(done))
    assert lint_docx(done)["findings"] == []
    assert TRUNCATION_MIN_DOC_CHARS >= 4_000  # boundary stays conservative


def test_docx_clean_document_zero_findings(tmp_path):
    d = docx_mod.Document()
    d.add_heading("Annual Report", level=1)
    d.add_paragraph("A complete introduction paragraph.")
    d.add_heading("Detail", level=2)
    d.add_paragraph("1. first item")
    d.add_paragraph("2. second item")
    t = d.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "Month"
    t.cell(1, 0).text = "Jan"
    t.cell(2, 0).text = "Feb"
    p = tmp_path / "clean.docx"
    d.save(str(p))
    res = lint_docx(p)
    assert res == {"findings": [], "ok": True}


def test_docx_unreadable_is_a_finding_never_a_raise(tmp_path):
    p = tmp_path / "corrupt.docx"
    p.write_bytes(b"not a zip archive at all")
    res = lint_docx(p)
    _check_shape(res)
    assert _codes(res) == ["docx-unreadable"] and not res["ok"]
    missing = lint_docx(tmp_path / "nope.docx")
    assert _codes(missing) == ["docx-unreadable"]


# ------------------------------------------------------------- xlsx lint -----


def _book(tmp_path: Path, name: str, rows: list, title: str = "Data") -> Path:
    wb = Workbook()
    ws = wb.active
    ws.title = title
    for row in rows:
        ws.append(row)
    p = tmp_path / name
    wb.save(str(p))
    return p


def test_xlsx_completely_empty_sheet(tmp_path):
    p = _book(tmp_path, "empty.xlsx", [])
    res = lint_xlsx(p)
    _check_shape(res)
    assert _codes(res) == ["xlsx-empty-sheet"] and not res["ok"]
    assert "completely empty" in res["findings"][0]["detail"]


def test_xlsx_header_only_sheet(tmp_path):
    p = _book(tmp_path, "header.xlsx", [["Name", "Total"]])
    res = lint_xlsx(p)
    _check_shape(res)
    assert _codes(res) == ["xlsx-empty-sheet"] and not res["ok"]
    assert "0 data rows" in res["findings"][0]["detail"]


def test_xlsx_empty_column_under_a_header(tmp_path):
    p = _book(
        tmp_path, "col.xlsx",
        [["Name", "Amount"], ["x", None], ["y", None]],
    )
    res = lint_xlsx(p)
    _check_shape(res)
    assert _codes(res) == ["xlsx-empty-column"] and not res["ok"]
    assert res["findings"][0]["where"] == "'Data'!B"
    assert "'Amount'" in res["findings"][0]["detail"]


def test_xlsx_error_literal_cells(tmp_path):
    p = _book(
        tmp_path, "err.xlsx",
        [["H1", "H2"], ["ok", "#REF!"], ["#DIV/0!", "fine"]],
    )
    res = lint_xlsx(p)
    _check_shape(res)
    assert _codes(res) == ["xlsx-error-cell", "xlsx-error-cell"]
    assert not res["ok"]
    wheres = {f["where"] for f in res["findings"]}
    assert wheres == {"'Data'!B2", "'Data'!A3"}


def test_xlsx_error_cells_are_capped_with_an_aggregate(tmp_path):
    rows = [["H"]] + [["#VALUE!"]] * 12
    p = _book(tmp_path, "flood.xlsx", rows)
    res = lint_xlsx(p)
    errs = [f for f in res["findings"] if f["code"] == "xlsx-error-cell"]
    assert len(errs) == 9  # 8 itemized + 1 honest aggregate, never a flood
    assert "4 more" in errs[-1]["detail"]


def test_xlsx_formula_beyond_used_range_conservatively(tmp_path):
    p = _book(
        tmp_path, "formulas.xlsx",
        [
            ["Label", "Val", "Calc"],
            ["x", 1, "=SUM(B2:B3)"],          # in range: clean
            ["y", 2, "=LOG10(B2)"],           # function name with digits: clean
            ["z", 3, "=SUM(B2:B100)"],        # beyond used range: flagged
            ["w", 4, "=SUM(Missing!A1:A99)"],  # sheet-qualified: skipped
        ],
    )
    res = lint_xlsx(p)
    _check_shape(res)
    assert _codes(res) == ["xlsx-formula-out-of-range"]
    assert res["ok"]  # warn severity
    assert "B2:B100" in res["findings"][0]["detail"]
    assert res["findings"][0]["where"] == "'Data'!C4"


def test_xlsx_number_as_text_in_majority_numeric_column(tmp_path):
    p = _book(
        tmp_path, "text.xlsx",
        [["Q", "Id"], [2, "007"], [3, "008"], ["4", "009"]],
    )
    res = lint_xlsx(p)
    _check_shape(res)
    # "4" flags (majority-numeric column Q); the leading-zero identifiers in
    # Id never do — that's the writer's own keep-as-text convention.
    assert _codes(res) == ["xlsx-number-as-text"]
    assert res["ok"]  # warn severity
    assert res["findings"][0]["where"] == "'Data'!A4"


def test_xlsx_chart_over_empty_range(tmp_path):
    from openpyxl.chart import BarChart, Reference

    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    for row in [["M", "S"], ["a", 1], ["b", 2]]:
        ws.append(row)
    ch = BarChart()
    ch.add_data(  # E2:E9 — inside no used cell, contains nothing
        Reference(ws, min_col=5, min_row=2, max_col=5, max_row=9),
        titles_from_data=False,
    )
    ws.add_chart(ch, "E12")
    p = tmp_path / "chart.xlsx"
    wb.save(str(p))
    res = lint_xlsx(p)
    _check_shape(res)
    assert _codes(res) == ["xlsx-chart-empty-range"] and not res["ok"]
    assert "chart 1" in res["findings"][0]["where"]


def test_xlsx_chart_over_real_data_is_clean(tmp_path):
    from openpyxl.chart import BarChart, Reference

    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    for row in [["M", "S"], ["a", 1], ["b", 2]]:
        ws.append(row)
    ch = BarChart()
    ch.add_data(
        Reference(ws, min_col=2, min_row=2, max_col=2, max_row=3),
        titles_from_data=False,
    )
    ws.add_chart(ch, "E2")
    p = tmp_path / "chart_ok.xlsx"
    wb.save(str(p))
    assert lint_xlsx(p) == {"findings": [], "ok": True}


def test_xlsx_clean_workbook_zero_findings(tmp_path):
    p = _book(
        tmp_path, "clean.xlsx",
        [
            ["Month", "Sales", "Total"],
            ["Jan", 100, "=SUM(B2:B2)"],
            ["Feb", 200, "=SUM(B2:B3)"],
        ],
    )
    assert lint_xlsx(p) == {"findings": [], "ok": True}


def test_xlsx_single_string_cell_is_a_note_not_a_header(tmp_path):
    # The writer itself puts one-line string content in a single cell — a lone
    # string cell is a note/title, never a "header row promising data". This
    # was a false positive on the app's own write_document output.
    from iron_jarvis.documents import write_document

    p = tmp_path / "note.xlsx"
    write_document(p, "a single note line")
    assert lint_xlsx(p) == {"findings": [], "ok": True}
    # ...while a MULTI-cell all-string single row is still the seeded defect.
    p2 = _book(tmp_path, "hdr.xlsx", [["Name", "Total"]])
    assert _codes(lint_xlsx(p2)) == ["xlsx-empty-sheet"]


def test_xlsx_unreadable_is_a_finding_never_a_raise(tmp_path):
    p = tmp_path / "corrupt.xlsx"
    p.write_bytes(b"definitely not a workbook")
    res = lint_xlsx(p)
    _check_shape(res)
    assert _codes(res) == ["xlsx-unreadable"] and not res["ok"]
    assert _codes(lint_xlsx(tmp_path / "gone.xlsx")) == ["xlsx-unreadable"]


def test_lint_document_dispatch(tmp_path):
    assert lint_document(tmp_path / "notes.txt") is None  # no linter != clean
    d = docx_mod.Document()
    d.add_paragraph("hello world")
    p = tmp_path / "d.docx"
    d.save(str(p))
    assert lint_document(p) == {"findings": [], "ok": True}
    # kind override mirrors write_document's dispatch (suffix lies, kind wins).
    from iron_jarvis.documents import write_document

    b = tmp_path / "report.bin"
    write_document(b, [["A", "B"], ["x", 1]], kind="xlsx")
    res = lint_document(b, kind="xlsx")
    assert res is not None and res["ok"]


# ------------------------------------------- batch refinement round ----------

_BAD_SHEETS = json.dumps({"sheets": {"Overview": [["Doc", "Total"]]}})  # header-only
_GOOD_SHEETS = json.dumps(
    {"sheets": {"Overview": [["Doc", "Total"], ["b.txt", 42]]}}
)


async def test_defective_xlsx_refined_and_failure_sheet_survives(tmp_path):
    src = tmp_path / "docs"
    _docs(src, a_txt="alpha body", b_txt="bravo body")
    out = tmp_path / "out"
    router = ScriptedRouter(
        ["junk one", "junk two",            # a.txt: extract + repair both fail
         _ext_reply("Bravo"),               # b.txt extracts
         _BAD_SHEETS,                       # synthesis: header-only -> lint error
         _GOOD_SHEETS]                      # the ONE refinement round fixes it
    )
    res = await run_batch(src, out, router, output="xlsx")
    assert res["deliverables"] and res["synthesis_errors"] == []
    qa = res["qa"]["synthesis.xlsx"]
    assert qa["refined"] is True and qa["refinement_error"] is None
    assert qa["findings"] == []  # findings describe the file ON DISK
    # The refinement prompt carried the lint findings + the previous reply.
    refine_user = router.calls[-1][1]
    assert "xlsx-empty-sheet" in refine_user and "Overview" in refine_user
    wb = load_workbook(out / "synthesis.xlsx")
    try:
        assert wb["Overview"]["A2"].value == "b.txt"  # refined content landed
        # The honesty sheet was RE-APPENDED BY CODE after the rewrite — the
        # model's refinement reply never mentioned it.
        assert "Not included" in wb.sheetnames
        rows = list(wb["Not included"].iter_rows(values_only=True))
        assert rows[0] == ("File", "Reason")
        assert rows[1][0] == "a.txt" and "extraction failed" in rows[1][1]
    finally:
        wb.close()


async def test_defective_docx_refined_and_failure_section_survives(tmp_path):
    src = tmp_path / "docs"
    _docs(src, a_txt="alpha body", b_txt="bravo body")
    out = tmp_path / "out"
    router = ScriptedRouter(
        ["junk one", "junk two",
         _ext_reply("Bravo"),
         "# Report\n\n## Empty",                 # empty section -> lint error
         "# Report\n\nAll good content here."]  # refinement fixes it
    )
    res = await run_batch(src, out, router, output="docx")
    qa = res["qa"]["synthesis.docx"]
    assert qa["refined"] is True and qa["findings"] == []
    text = extract_text(out / "synthesis.docx")
    assert "All good content here" in text
    assert "Documents not included" in text and "a.txt" in text


async def test_refinement_garbage_keeps_original_and_records_error(tmp_path):
    src = tmp_path / "docs"
    _docs(src, a_txt="alpha body")
    out = tmp_path / "out"
    router = ScriptedRouter(
        [_ext_reply("Alpha"), _BAD_SHEETS, "utterly not json"]
    )
    res = await run_batch(src, out, router, output="xlsx")
    # The deliverable stays SHIPPED — QA never un-writes a produced file.
    assert res["deliverables"] and res["synthesis_errors"] == []
    qa = res["qa"]["synthesis.xlsx"]
    assert qa["refined"] is False
    assert qa["refinement_error"] and "ValueError" in qa["refinement_error"]
    assert "xlsx-empty-sheet" in [f["code"] for f in qa["findings"]]
    wb = load_workbook(out / "synthesis.xlsx")
    try:  # the original (header-only) file is untouched
        assert list(wb["Overview"].iter_rows(values_only=True)) == [("Doc", "Total")]
    finally:
        wb.close()


async def test_refinement_provider_failure_recorded_not_raised(tmp_path):
    src = tmp_path / "docs"
    _docs(src, a_txt="alpha body")
    out = tmp_path / "out"
    router = ScriptedRouter(
        [_ext_reply("Alpha"), _BAD_SHEETS, RuntimeError("provider down")]
    )
    res = await run_batch(src, out, router, output="xlsx")
    assert res["deliverables"]  # still shipped
    qa = res["qa"]["synthesis.xlsx"]
    assert qa["refined"] is False and "provider down" in qa["refinement_error"]


async def test_refinement_that_lints_worse_keeps_original(tmp_path):
    src = tmp_path / "docs"
    _docs(src, a_txt="alpha body")
    out = tmp_path / "out"
    router = ScriptedRouter(
        [_ext_reply("Alpha"),
         "# Report\n\nIntro paragraph.\n\n## Gap",  # 1 error (empty section)
         "## A\n\n## B"]                            # refinement: 2 errors — WORSE
    )
    res = await run_batch(src, out, router, output="docx")
    qa = res["qa"]["synthesis.docx"]
    assert qa["refined"] is False
    assert "did not lint better" in qa["refinement_error"]
    assert [f["code"] for f in qa["findings"]] == ["docx-empty-section"]
    text = extract_text(out / "synthesis.docx")
    assert "Intro paragraph" in text  # the original is what shipped
    # ...and the rejected candidate temp file was cleaned up.
    assert not list(out.glob(".synthesis.refine-*"))


async def test_refinement_tie_keeps_original(tmp_path):
    # STRICTLY better or nothing: an equal-rank refinement is churn, not gain.
    src = tmp_path / "docs"
    _docs(src, a_txt="alpha body")
    out = tmp_path / "out"
    router = ScriptedRouter(
        [_ext_reply("Alpha"),
         "# Report\n\nIntro paragraph.\n\n## Gap",       # 1 error
         "# Report\n\nIntro paragraph.\n\n## StillGap"]  # also 1 error — a tie
    )
    res = await run_batch(src, out, router, output="docx")
    qa = res["qa"]["synthesis.docx"]
    assert qa["refined"] is False and "did not lint better" in qa["refinement_error"]
    text = extract_text(out / "synthesis.docx")
    assert "Gap" in text and "StillGap" not in text


async def test_warnings_alone_never_spend_a_refinement_round(tmp_path):
    src = tmp_path / "docs"
    _docs(src, a_txt="alpha body")
    out = tmp_path / "out"
    big_md = (
        "# Report\n\n"
        + ("Solid sentence content here. " * 250).strip()
        + "\n\nThe final thought was left dangling and contin"
    )
    # Exactly 2 scripted replies: if the warn triggered refinement, the
    # ScriptedRouter's over-call assertion would fail this test.
    router = ScriptedRouter([_ext_reply("Alpha"), big_md])
    res = await run_batch(src, out, router, output="docx")
    qa = res["qa"]["synthesis.docx"]
    assert [f["code"] for f in qa["findings"]] == ["docx-truncated-tail"]
    assert qa["refined"] is False and qa["refinement_error"] is None


async def test_clean_deliverable_gets_a_clean_qa_entry(tmp_path):
    src = tmp_path / "docs"
    _docs(src, a_txt="alpha body")
    out = tmp_path / "out"
    router = ScriptedRouter(
        [_ext_reply("Alpha"), "# Batch Report\n\n- combined finding one"]
    )
    res = await run_batch(src, out, router, output="docx")
    assert res["qa"] == {
        "synthesis.docx": {
            "findings": [], "refined": False, "refinement_error": None,
        }
    }


async def test_refinement_candidate_lint_crash_keeps_original_no_debris(
    tmp_path, monkeypatch
):
    # Even if linting the CANDIDATE blows up, the shipped original stays on
    # disk, the failure is recorded, and no .refine-* temp file is left behind.
    from iron_jarvis.documents import write_document
    from iron_jarvis.documents.lint import lint_xlsx as real_lint

    target = tmp_path / "synthesis.xlsx"
    write_document(target, {"sheets": {"O": [["Doc", "Total"]]}})
    findings = real_lint(target)["findings"]
    assert findings

    def boom(path):
        raise RuntimeError("lint crashed on the candidate")

    monkeypatch.setitem(batch_mod._LINT_FOR_FMT, "xlsx", boom)
    router = ScriptedRouter([_GOOD_SHEETS])
    refined, err, kept = await refine_deliverable(
        target, "xlsx", {}, findings, router=router, user="u", excluded=[],
    )
    assert refined is False and "lint crashed" in err and kept == findings
    assert not list(tmp_path.glob(".*refine-*"))
    wb = load_workbook(target)
    try:  # the original is byte-honest: still the header-only sheet
        assert list(wb["O"].iter_rows(values_only=True)) == [("Doc", "Total")]
    finally:
        wb.close()


async def test_refinement_threads_the_synthesize_role_llm(tmp_path):
    # v1.135.0 interplay: the refinement one-shot rides the SAME resolved
    # synthesize role as first-pass synthesis (provider/model reach the router).
    from iron_jarvis.documents import write_document
    from iron_jarvis.documents.lint import lint_xlsx as real_lint

    class Applied:
        applied = True
        provider = "prov-x"
        model = "model-y"

    captured = []

    class KwRouter(ScriptedRouter):
        async def complete(self, *, system, messages, tools, task_class=None, **kw):
            captured.append((task_class, kw.get("provider"), kw.get("model")))
            return await super().complete(
                system=system, messages=messages, tools=tools,
                task_class=task_class,
            )

    target = tmp_path / "synthesis.xlsx"
    write_document(target, {"sheets": {"O": [["Doc", "Total"]]}})
    router = KwRouter([_GOOD_SHEETS])
    refined, err, kept = await refine_deliverable(
        target, "xlsx", {}, real_lint(target)["findings"],
        router=router, user="u", excluded=[], llm=Applied(),
    )
    assert refined is True and err is None and kept == []
    assert captured == [("synthesize", "prov-x", "model-y")]


# ------------------------------------------------------- tool surfaces -------


async def test_batch_tool_surfaces_qa(tmp_path):
    src = tmp_path / "docs"
    _docs(src, a_txt="alpha body")
    router = ScriptedRouter([_ext_reply("Alpha"), _BAD_SHEETS, _GOOD_SHEETS])
    tool = next(
        t for t in document_tools(router_resolver=lambda: router)
        if t.name == "batch_documents"
    )
    res = await tool.execute(
        {"folder": str(src), "output": "xlsx"}, _ctx(tmp_path / "ws")
    )
    assert res.ok, res.error
    assert res.data["qa"]["synthesis.xlsx"]["refined"] is True
    assert "qa synthesis.xlsx: clean after one refinement round" in res.output


async def test_write_document_tool_surfaces_findings_no_refinement(tmp_path):
    tool = next(t for t in document_tools() if t.name == "write_document")
    ctx = _ctx(tmp_path / "ws")
    res = await tool.execute(
        {"path": "report.docx", "content": "# Title\n\nBody text.\n\n## Empty"},
        ctx,
    )
    assert res.ok, res.error  # findings never fail a landed write
    assert "qa lint: 1 error(s), 0 warning(s)" in res.output
    assert "docx-empty-section" in res.output
    assert res.data["lint"]["ok"] is False
    assert [f["code"] for f in res.data["lint"]["findings"]] == ["docx-empty-section"]
    # The tool NEVER refines — the calling agent is the refiner. The file on
    # disk is exactly what was asked for, defect and all.
    assert "Empty" in extract_text(tmp_path / "ws" / "report.docx")


async def test_write_document_tool_clean_line_and_no_lint_for_text(tmp_path):
    tool = next(t for t in document_tools() if t.name == "write_document")
    ctx = _ctx(tmp_path / "ws")
    res = await tool.execute(
        {
            "path": "clean.xlsx",
            "content": {"sheets": {"D": [["Month", "Sales"], ["Jan", 100]]}},
        },
        ctx,
    )
    assert res.ok, res.error
    assert "qa lint: clean" in res.output
    assert res.data["lint"] == {"findings": [], "ok": True}

    res = await tool.execute({"path": "notes.txt", "content": "hello"}, ctx)
    assert res.ok, res.error
    assert "qa lint" not in res.output  # no linter for .txt — and no fake "clean"
    assert "lint" not in res.data
