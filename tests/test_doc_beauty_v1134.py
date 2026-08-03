"""Document beauty layer (v1.134.0) — declarative themes, charts, and polish.

The engine owns beauty: models pick OPTIONS (theme name, cover, footer,
banded, charts) instead of writing styling code. These tests introspect the
real outputs (python-docx / openpyxl / raw zip parts) to prove the options
land — and, just as load-bearing, that the NO-options path still produces
exactly the legacy structure (styles, no cover, no fills, no charts).
"""

from __future__ import annotations

import zipfile
from datetime import datetime
from pathlib import Path

import docx as docx_mod
from openpyxl import Workbook, load_workbook

from iron_jarvis.documents import write_document
from iron_jarvis.documents.themes import THEME_NAMES, THEMES, get_theme
from iron_jarvis.documents.tools import WriteDocumentTool, document_tools
from iron_jarvis.tools.base import ToolContext

_MD = "# Annual Report\n\nSome intro text.\n\n## Detail\n\n- one\n- two\n"

_SHEET_ROWS = [
    ["Month", "Sales", "Notes"],
    ["Jan", "100", "x" * 80],  # 80-char cell exercises the autosize clamp
    ["Feb", "200", "ok"],
    ["Mar", "300", "ok"],
]


def _ctx(ws: Path) -> ToolContext:
    return ToolContext(
        workspace=ws, session_id="t", agent_run_id="t",
        config=None, event_bus=None, engine=None,
    )


def _zip_chart_xml(path: Path) -> str:
    """Concatenated chart-part XML — '' when the workbook has no charts."""
    with zipfile.ZipFile(path) as z:
        parts = [n for n in z.namelist() if n.startswith("xl/charts/chart")]
        return "".join(z.read(n).decode("utf-8", "replace") for n in parts)


# --- themes module -------------------------------------------------------------


def test_theme_registry_shape():
    assert set(THEME_NAMES) == {"professional", "minimal", "warm"}
    for t in THEMES.values():
        assert t.heading_font and t.body_font
        assert len(t.accent_rgb) == 3
        assert len(t.table_header_fill) == 6 and len(t.band_fill) == 6
        assert len(t.margins_in) == 4
    assert get_theme("PROFESSIONAL") is THEMES["professional"]
    assert get_theme("nope") is None and get_theme(None) is None


# --- docx: theme / cover / header / footer -------------------------------------


def test_docx_theme_lands_in_named_styles(tmp_path):
    p = tmp_path / "r.docx"
    write_document(p, _MD, options={"theme": "professional"})
    d = docx_mod.Document(str(p))
    t = THEMES["professional"]
    h1 = d.styles["Heading 1"]
    assert h1.font.name == t.heading_font
    assert h1.font.size.pt == t.heading_size(1)
    assert str(h1.font.color.rgb) == ("%02X%02X%02X" % t.accent_rgb)
    normal = d.styles["Normal"]
    assert normal.font.name == t.body_font
    assert normal.font.size.pt == t.body_size_pt


def test_docx_theme_margins(tmp_path):
    p = tmp_path / "m.docx"
    write_document(p, _MD, options={"theme": "minimal"})
    sec = docx_mod.Document(str(p)).sections[0]
    assert sec.top_margin.inches == 0.75
    assert sec.left_margin.inches == 0.75


def test_docx_cover_page_present_when_asked(tmp_path):
    p = tmp_path / "c.docx"
    write_document(
        p, _MD,
        options={"theme": "professional", "cover": True, "subtitle": "Q3 Review"},
    )
    d = docx_mod.Document(str(p))
    texts = [para.text for para in d.paragraphs]
    # Title from the first H1, subtitle, and today's date on the cover.
    assert "Annual Report" in texts
    assert "Q3 Review" in texts
    assert any(str(datetime.now().year) in x for x in texts)
    # The cover ends in a real page break; the H1 still opens the body.
    assert 'w:type="page"' in d.element.xml
    assert any(
        para.style.name == "Heading 1" and para.text == "Annual Report"
        for para in d.paragraphs
    )


def test_docx_footer_page_number_field_and_header(tmp_path):
    p = tmp_path / "f.docx"
    write_document(
        p, _MD,
        options={"header_text": "RealDealCPA", "footer": "page-numbers"},
    )
    sec = docx_mod.Document(str(p)).sections[0]
    assert sec.header.paragraphs[0].text == "RealDealCPA"
    fxml = sec.footer.paragraphs[0]._p.xml
    # The standard fieldcode dance: fldChar begin/end around instrText PAGE.
    assert "fldChar" in fxml and "PAGE" in fxml


def test_docx_footer_text_with_page_token(tmp_path):
    p = tmp_path / "ft.docx"
    write_document(p, _MD, options={"footer": "Report — page {page}"})
    fp = docx_mod.Document(str(p)).sections[0].footer.paragraphs[0]
    fxml = fp._p.xml
    assert "Report — page " in fp.text
    assert "{page}" not in fxml  # the token became the live field
    assert "fldChar" in fxml and "PAGE" in fxml


def test_docx_no_options_is_legacy(tmp_path):
    """LEGACY PIN: without options nothing is styled — styles match a fresh
    python-docx default document, no cover break, empty header/footer."""
    p = tmp_path / "plain.docx"
    write_document(p, _MD)
    d = docx_mod.Document(str(p))
    fresh = docx_mod.Document()
    for style in ("Normal", "Heading 1", "Heading 2"):
        assert d.styles[style].font.name == fresh.styles[style].font.name
        assert d.styles[style].font.size == fresh.styles[style].font.size
    assert 'w:type="page"' not in d.element.xml  # no cover page break
    sec = d.sections[0]
    assert sec.header.paragraphs[0].text == ""
    assert "fldChar" not in sec.footer.paragraphs[0]._p.xml
    assert sec.top_margin == fresh.sections[0].top_margin
    assert any(para.style.name == "Heading 1" for para in d.paragraphs)


def test_docx_unknown_theme_warns_but_writes(tmp_path):
    p = tmp_path / "u.docx"
    warns: list[str] = []
    write_document(p, _MD, options={"theme": "neon"}, warnings=warns)
    assert p.is_file()
    assert any("unknown theme" in w and "neon" in w for w in warns)


def test_options_on_other_formats_warn_not_crash(tmp_path):
    p = tmp_path / "r.pdf"
    warns: list[str] = []
    write_document(p, _MD, options={"theme": "professional"}, warnings=warns)
    assert p.is_file() and p.stat().st_size > 0
    assert any(".docx/.xlsx" in w for w in warns)


# --- xlsx: beauty options + charts ---------------------------------------------


def test_xlsx_beauty_options_land(tmp_path):
    p = tmp_path / "b.xlsx"
    write_document(
        p,
        {"sheets": {"Data": _SHEET_ROWS}},
        options={
            "theme": "professional",
            "banded": True,
            "freeze_header": True,
            "number_formats": {"B": "#,##0.00"},
        },
    )
    ws = load_workbook(str(p))["Data"]
    t = THEMES["professional"]
    # Themed header row: fill + contrast font + bold.
    assert ws["A1"].fill.start_color.rgb.endswith(t.table_header_fill)
    assert ws["A1"].font.bold
    assert str(ws["A1"].font.color.rgb).endswith(t.table_header_font)
    # Banded: row 3 shaded, row 2 (first data row) clean.
    assert ws["A3"].fill.start_color.rgb.endswith(t.band_fill)
    assert ws["A2"].fill.patternType is None
    assert ws.freeze_panes == "A2"
    # Autosize (default on with a theme): widths within the 8..60 clamp.
    assert ws.column_dimensions["A"].width == 8  # short labels hit the floor
    assert ws.column_dimensions["C"].width == 60  # 80-char note hits the cap
    # Number formats on data rows only; header label keeps General.
    assert ws["B2"].number_format == "#,##0.00"
    assert ws["B4"].number_format == "#,##0.00"
    assert ws["B1"].number_format == "General"


def test_xlsx_plain_rows_with_options(tmp_path):
    # The non-sheets (list[list]) path gets the same beauty treatment.
    p = tmp_path / "flat.xlsx"
    write_document(
        p, _SHEET_ROWS,
        options={"theme": "minimal", "freeze_header": True},
    )
    ws = load_workbook(str(p)).active
    assert ws.freeze_panes == "A2"
    assert 8 <= ws.column_dimensions["A"].width <= 60


def test_xlsx_chart_written_with_right_ranges(tmp_path):
    p = tmp_path / "chart.xlsx"
    write_document(
        p,
        {"sheets": {"Data": {
            "rows": _SHEET_ROWS,
            "charts": [{
                "type": "bar", "title": "Sales by Month",
                "data_range": "B2:B4", "categories_range": "A2:A4",
                "anchor": "E2",
            }],
        }}},
        options={"theme": "professional"},
    )
    xml = _zip_chart_xml(p)
    assert "barChart" in xml
    assert "$B$2:$B$4" in xml  # data reference
    assert "$A$2:$A$4" in xml  # categories reference
    assert "Sales by Month" in xml


def test_xlsx_charts_work_without_options(tmp_path):
    # Charts are a CONTENT extension — no options dict required.
    p = tmp_path / "co.xlsx"
    write_document(p, {"sheets": {"D": {
        "rows": _SHEET_ROWS,
        "charts": [{"type": "line", "data_range": "B2:B4"}],
    }}})
    assert "lineChart" in _zip_chart_xml(p)


def test_xlsx_invalid_chart_skipped_with_warning(tmp_path):
    p = tmp_path / "bad.xlsx"
    warns: list[str] = []
    write_document(
        p,
        {"sheets": {"Data": {
            "rows": _SHEET_ROWS,
            "charts": [
                {"type": "bar", "data_range": "ZZ2:ZZ400"},  # outside used cells
                {"type": "donut", "data_range": "B2:B4"},    # unknown type
                {"type": "pie", "data_range": "not-a-range"},
            ],
        }}},
        warnings=warns,
    )
    assert p.is_file()  # the write always lands
    assert _zip_chart_xml(p) == ""  # nothing invalid rendered
    assert len(warns) == 3
    assert any("ZZ2:ZZ400" in w for w in warns)
    assert any("donut" in w for w in warns)
    # And the sheet data itself is intact.
    assert load_workbook(str(p))["Data"]["B2"].value == 100


def test_xlsx_no_options_is_legacy(tmp_path):
    """LEGACY PIN: without options the workbook keeps the historical shape —
    bold+frozen header, legacy widths, no fills, no charts."""
    p = tmp_path / "legacy.xlsx"
    write_document(p, {"sheets": {"Data": _SHEET_ROWS}})
    ws = load_workbook(str(p))["Data"]
    assert ws["A1"].font.bold  # historical header bolding
    assert ws.freeze_panes == "A2"  # historical header freeze
    assert ws["A1"].fill.patternType is None  # NO themed header fill
    assert ws["A3"].fill.patternType is None  # NO banding
    assert _zip_chart_xml(p) == ""
    # Historical width formula: min(max(len + 2, 8), 60).
    assert ws.column_dimensions["A"].width == 8
    assert ws.column_dimensions["C"].width == 60
    assert ws["B2"].number_format == "General"


# --- write_document tool surface ------------------------------------------------


def test_tool_theme_enum_matches_registry():
    enum = WriteDocumentTool.input_schema["properties"]["theme"]["enum"]
    assert enum == list(THEME_NAMES)


async def test_tool_threads_theme_and_surfaces_warnings(tmp_path):
    ws_dir = tmp_path / "ws"
    ws_dir.mkdir()
    tool = next(t for t in document_tools() if t.name == "write_document")
    res = await tool.execute(
        {
            "path": "report.docx",
            "content": _MD,
            "theme": "warm",
            "options": {"cover": True},
        },
        _ctx(ws_dir),
    )
    assert res.ok, res.error
    d = docx_mod.Document(str(ws_dir / "report.docx"))
    assert d.styles["Heading 1"].font.name == THEMES["warm"].heading_font
    assert 'w:type="page"' in d.element.xml

    res = await tool.execute(
        {
            "path": "s.xlsx",
            "content": {"sheets": {"D": {
                "rows": _SHEET_ROWS,
                "charts": [{"type": "bar", "data_range": "Q1:Q9"}],
            }}},
            "theme": "professional",
        },
        _ctx(ws_dir),
    )
    assert res.ok, res.error
    assert "warning:" in res.output  # the skipped chart is surfaced
    assert res.data["warnings"]


async def test_tool_without_theme_is_legacy(tmp_path):
    ws_dir = tmp_path / "ws"
    ws_dir.mkdir()
    tool = next(t for t in document_tools() if t.name == "write_document")
    res = await tool.execute(
        {"path": "plain.docx", "content": _MD}, _ctx(ws_dir)
    )
    assert res.ok, res.error
    assert res.output.startswith("wrote ")
    assert "warnings" not in res.data
    d = docx_mod.Document(str(ws_dir / "plain.docx"))
    assert d.styles["Heading 1"].font.name == docx_mod.Document().styles["Heading 1"].font.name


# --- excel_apply_spec: beauty in specs + legacy specs unchanged -----------------


def _seed_book(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    for row in [["Month", "Sales"], ["Jan", 100], ["Feb", 200], ["Mar", 300]]:
        ws.append(row)
    wb.save(str(path))


_LEGACY_SPEC = {
    "sheet": "Data",
    "cells": {
        "B5": {"formula": "=SUM(B2:B4)", "bold": True, "number_format": "#,##0"},
    },
    "column_widths": {"A": 14.0},
    "merges": [],
}


async def test_apply_spec_legacy_spec_unchanged(tmp_path):
    ws_dir = tmp_path / "ws"
    ws_dir.mkdir()
    _seed_book(ws_dir / "book.xlsx")
    tool = next(t for t in document_tools() if t.name == "excel_apply_spec")
    res = await tool.execute(
        {"path": "book.xlsx", "spec": dict(_LEGACY_SPEC)}, _ctx(ws_dir)
    )
    assert res.ok, res.error
    # Output line reads exactly as it always has — no beauty chatter.
    assert res.output.startswith("reproduced + validated")
    assert "chart" not in res.output and "warning" not in res.output
    assert res.data["charts_added"] == 0
    assert res.data["beauty_warnings"] == []
    ws = load_workbook(str(ws_dir / "book.xlsx"))["Data"]
    assert ws["B5"].value == "=SUM(B2:B4)"
    assert ws["B5"].font.bold
    assert ws.column_dimensions["A"].width == 14.0
    # And nothing beautified behind the spec's back.
    assert ws["A1"].fill.patternType is None
    assert ws.freeze_panes is None
    assert _zip_chart_xml(ws_dir / "book.xlsx") == ""


async def test_apply_spec_with_beauty_and_charts(tmp_path):
    ws_dir = tmp_path / "ws"
    ws_dir.mkdir()
    _seed_book(ws_dir / "book.xlsx")
    spec = {
        **_LEGACY_SPEC,
        "options": {"theme": "professional", "banded": True, "freeze_header": True},
        "charts": [
            {"type": "pie", "title": "Mix", "data_range": "B2:B4",
             "categories_range": "A2:A4", "anchor": "E2"},
            {"type": "bar", "data_range": "ZZ1:ZZ99"},  # invalid -> warning
        ],
    }
    tool = next(t for t in document_tools() if t.name == "excel_apply_spec")
    res = await tool.execute({"path": "book.xlsx", "spec": spec}, _ctx(ws_dir))
    assert res.ok, res.error
    assert res.data["charts_added"] == 1
    assert any("ZZ1:ZZ99" in w for w in res.data["beauty_warnings"])
    assert "warning:" in res.output and "+ 1 chart(s)" in res.output
    ws = load_workbook(str(ws_dir / "book.xlsx"))["Data"]
    t = THEMES["professional"]
    assert ws["A1"].fill.start_color.rgb.endswith(t.table_header_fill)
    assert ws["A3"].fill.start_color.rgb.endswith(t.band_fill)
    assert ws.freeze_panes == "A2"
    xml = _zip_chart_xml(ws_dir / "book.xlsx")
    assert "pieChart" in xml and "$B$2:$B$4" in xml
